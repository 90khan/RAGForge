from pathlib import Path

from docx import Document as DocxDocument

from loaders.base_loader import BaseLoader
from models.document import Document


class DOCXLoader(BaseLoader):

    def load(
        self,
        file_path: Path
    ) -> Document:

        doc = DocxDocument(file_path)

        text = "\n".join(
            p.text
            for p in doc.paragraphs
        )

        return Document(
            id=file_path.stem,
            source=file_path,
            text=text,
        )